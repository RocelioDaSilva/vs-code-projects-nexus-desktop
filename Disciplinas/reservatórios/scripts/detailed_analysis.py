#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Comprehensive analysis of colleague documents
Extracts key sections for comparison with main document
"""

import os
from pathlib import Path
from PyPDF2 import PdfReader
from docx import Document

colleague_dir = Path("c:/Users/PCGAME/Desktop/reservatórios/histo/referências de meus colegas")

# ============ DOCX ANALYSIS ============
print("\n" + "="*100)
print("DOCUMENTO (3) (2) (1).docx - COMPREHENSIVE ANALYSIS")
print("="*100)

docx_file_1 = colleague_dir / "Documento (3) (2) (1).docx"
try:
    doc1 = Document(docx_file_1)
    all_text_1 = [para.text for para in doc1.paragraphs if para.text.strip()]
    
    print(f"\nTotal Paragraphs: {len(all_text_1)}")
    print("\nFull Content (last 80 paragraphs):")
    for i, line in enumerate(all_text_1[-80:], 1):
        print(f"{i:3d}: {line[:140]}")
        
except Exception as e:
    print(f"ERROR: {e}")

# ============ DOCX 2 ANALYSIS ============
print("\n\n" + "="*100)
print("DOCUMENTO (3) (3) (1).docx - COMPREHENSIVE ANALYSIS")
print("="*100)

docx_file_2 = colleague_dir / "Documento (3) (3) (1).docx"
try:
    doc2 = Document(docx_file_2)
    all_text_2 = [para.text for para in doc2.paragraphs if para.text.strip()]
    
    print(f"\nTotal Paragraphs: {len(all_text_2)}")
    print("\nFull Content:")
    for i, line in enumerate(all_text_2, 1):
        print(f"{i:3d}: {line[:140]}")
        
except Exception as e:
    print(f"ERROR: {e}")
